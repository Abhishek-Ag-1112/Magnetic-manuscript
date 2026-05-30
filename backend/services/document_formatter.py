"""
Document Formatter Service — Generates PRODUCTION-QUALITY, publish-ready DOCX.
Applies journal-specific margins, fonts, spacing, heading styles, proper reference
formatting, page numbers, professional layout, two-column support, and LaTeX export.
"""
import os
import re
import subprocess
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

from utils.helpers import OUTPUT_DIR


# ─── ROMAN NUMERAL CONVERSION ───

ROMAN_NUMERALS = [
    (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
    (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
    (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
]


def _to_roman(num: int) -> str:
    """Convert integer to Roman numeral string."""
    result = []
    for value, numeral in ROMAN_NUMERALS:
        while num >= value:
            result.append(numeral)
            num -= value
    return ''.join(result)


def format_document(structured_json: dict, journal_rules: dict, session_id: str) -> dict:
    """Generate a publish-ready DOCX file from structured content."""
    if Document is None:
        raise ImportError("python-docx is required for document formatting")

    doc = Document()

    # Set up document styles and page layout
    _setup_document_styles(doc, journal_rules)
    _apply_page_layout(doc, journal_rules)

    font_name = journal_rules.get("font", "Times New Roman")
    font_size = journal_rules.get("font_size", 12)
    line_spacing_val = journal_rules.get("line_spacing", 1.5)
    heading_style = journal_rules.get("heading_style", "bold_title_case")
    is_numbered = journal_rules.get("section_numbered", False)

    # ── TITLE ──
    title = structured_json.get("title", "Untitled Manuscript")
    if title:
        title_para = doc.add_paragraph()
        title_align = journal_rules.get("title_alignment", "center")
        if title_align == "left":
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(6)
        title_para.space_before = Pt(24)
        title_run = title_para.add_run(title)
        title_run.bold = journal_rules.get("title_bold", True)
        title_font_sz = journal_rules.get("title_font_size", font_size + 4)
        title_run.font.size = Pt(title_font_sz)
        title_run.font.name = font_name
        _set_run_font(title_run, font_name)

    # ── AUTHORS ──
    authors = structured_json.get("authors", [])
    if authors:
        authors_text = ", ".join(authors) if isinstance(authors, list) else str(authors)
        authors_para = doc.add_paragraph()
        author_align = journal_rules.get("author_alignment", "center")
        if author_align == "left":
            authors_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_para.space_after = Pt(4)
        authors_run = authors_para.add_run(authors_text)
        authors_run.font.size = Pt(font_size)
        authors_run.font.name = font_name
        _set_run_font(authors_run, font_name)

    # ── DIVIDER AFTER TITLE BLOCK ──
    _add_divider(doc)

    # ── ABSTRACT ──
    abstract = structured_json.get("abstract", "")
    abstract_style = journal_rules.get("abstract_style", "italic_indented")

    if abstract:
        if abstract_style == "bold_label":
            # IEEE-style: Bold "Abstract—" followed by text on same paragraph
            abstract_para = doc.add_paragraph()
            abstract_para.space_before = Pt(10)
            abstract_para.space_after = Pt(8)
            _apply_line_spacing(abstract_para, line_spacing_val)

            label_run = abstract_para.add_run("Abstract—")
            label_run.bold = True
            label_run.italic = True
            label_run.font.size = Pt(font_size)
            label_run.font.name = font_name
            _set_run_font(label_run, font_name)

            text_run = abstract_para.add_run(abstract)
            text_run.font.size = Pt(font_size)
            text_run.font.name = font_name
            _set_run_font(text_run, font_name)

        elif abstract_style == "italic_indented":
            # Nature-style: No heading, italic text with indent
            abstract_para = doc.add_paragraph()
            abstract_para.paragraph_format.first_line_indent = Inches(0)
            abstract_para.paragraph_format.left_indent = Inches(0.5)
            abstract_para.paragraph_format.right_indent = Inches(0.5)
            abstract_para.space_before = Pt(10)
            abstract_para.space_after = Pt(8)
            _apply_line_spacing(abstract_para, line_spacing_val)

            abstract_run = abstract_para.add_run(abstract)
            abstract_run.font.size = Pt(font_size)
            abstract_run.font.name = font_name
            abstract_run.italic = True
            _set_run_font(abstract_run, font_name)

        else:
            # Default: Labeled heading + indented italic text
            _add_section_heading(doc, "Abstract", journal_rules, level=1)

            abstract_para = doc.add_paragraph()
            abstract_para.paragraph_format.first_line_indent = Inches(0)
            abstract_para.paragraph_format.left_indent = Inches(0.5)
            abstract_para.paragraph_format.right_indent = Inches(0.5)
            _apply_line_spacing(abstract_para, line_spacing_val)

            abstract_run = abstract_para.add_run(abstract)
            abstract_run.font.size = Pt(font_size)
            abstract_run.font.name = font_name
            abstract_run.italic = True
            _set_run_font(abstract_run, font_name)

    # ── KEYWORDS ──
    keywords = structured_json.get("keywords", [])
    if keywords:
        kw_para = doc.add_paragraph()
        if abstract_style == "bold_label":
            # IEEE-style: Index Terms
            kw_label = kw_para.add_run("Index Terms—")
            kw_label.bold = True
            kw_label.italic = True
        else:
            kw_para.paragraph_format.left_indent = Inches(0.5)
            kw_para.paragraph_format.right_indent = Inches(0.5)
            kw_label = kw_para.add_run("Keywords: ")
            kw_label.bold = True

        kw_label.font.size = Pt(font_size)
        kw_label.font.name = font_name
        _set_run_font(kw_label, font_name)

        kw_text = ", ".join(keywords) if isinstance(keywords, list) else str(keywords)
        kw_run = kw_para.add_run(kw_text)
        kw_run.font.size = Pt(font_size)
        kw_run.font.name = font_name
        kw_run.italic = True
        _set_run_font(kw_run, font_name)

        kw_para.space_before = Pt(6)
        kw_para.space_after = Pt(12)

    # ── DIVIDER BEFORE MAIN BODY ──
    _add_divider(doc)

    # ── TWO-COLUMN LAYOUT (IEEE/ACM) ──
    columns = journal_rules.get("columns", 1)
    if columns == 2:
        _apply_two_column_layout(doc)

    # ── BODY SECTIONS ──
    section_order = journal_rules.get("section_order", [])
    sections = structured_json.get("sections", [])

    # Build mapping for ordered output
    section_map = {}
    for sec in sections:
        heading = sec.get("heading", "").strip()
        key = _normalize_section_key(heading)
        if key not in section_map:
            section_map[key] = sec

    # Add sections in journal order
    added_keys = set()
    section_counter = 1
    for ordered_section in section_order:
        key = _normalize_section_key(ordered_section)
        if key in ("abstract", "references", "keywords"):
            continue

        sec = section_map.get(key)
        if sec:
            _add_full_section(doc, sec, journal_rules, line_spacing_val,
                              section_num=section_counter)
            section_counter += 1
            added_keys.add(key)

    # Add remaining sections not in the specified order
    for sec in sections:
        heading = sec.get("heading", "").strip()
        key = _normalize_section_key(heading)
        if key not in added_keys and key not in ("abstract", "references", "keywords"):
            _add_full_section(doc, sec, journal_rules, line_spacing_val,
                              section_num=section_counter)
            section_counter += 1

    # ── TABLES (with auto-numbering) ──
    extracted_tables = structured_json.get("extracted_tables", [])
    if extracted_tables:
        _add_divider(doc)
        for idx, tbl in enumerate(extracted_tables, start=1):
            _add_formatted_table(doc, tbl, journal_rules, table_number=idx)

    # ── IMAGES / FIGURES (with auto-numbering) ──
    extracted_images = structured_json.get("extracted_images", [])
    if extracted_images:
        _add_divider(doc)
        for idx, img in enumerate(extracted_images, start=1):
            _add_image_to_doc(doc, img, journal_rules, figure_number=idx)

    # ── REFERENCES ──
    references = structured_json.get("references", [])
    reference_style = journal_rules.get("reference_style", "numbered_brackets")
    if references:
        _add_divider(doc)
        _add_section_heading(doc, "References", journal_rules, level=1)

        ref_font_size = journal_rules.get("reference_font_size", max(font_size - 1, 9))
        ref_hanging = journal_rules.get("reference_hanging_indent", 0.35)

        for i, ref in enumerate(references):
            ref_text = ref if isinstance(ref, str) else str(ref)
            ref_para = doc.add_paragraph()

            if reference_style == "numbered_superscript":
                # Nature-style: superscript number
                ref_num_run = ref_para.add_run(f"{i+1}. ")
                ref_num_run.font.size = Pt(ref_font_size)
                ref_num_run.font.name = font_name
                _set_run_font(ref_num_run, font_name)
            elif reference_style == "numbered_brackets":
                # IEEE-style: [1]
                ref_num_run = ref_para.add_run(f"[{i+1}]  ")
                ref_num_run.bold = True
                ref_num_run.font.size = Pt(ref_font_size)
                ref_num_run.font.name = font_name
                _set_run_font(ref_num_run, font_name)
            else:
                # Default numbered
                ref_num_run = ref_para.add_run(f"[{i+1}]  ")
                ref_num_run.font.size = Pt(ref_font_size)
                ref_num_run.font.name = font_name
                _set_run_font(ref_num_run, font_name)

            # Add DOI hyperlink if present
            doi_match = re.search(r'(10\.\d{4,}/[^\s,]+)', ref_text)
            if doi_match:
                doi = doi_match.group(1)
                text_before_doi = ref_text[:doi_match.start()].strip()
                text_after_doi = ref_text[doi_match.end():].strip()

                if text_before_doi:
                    ref_run = ref_para.add_run(text_before_doi + " ")
                    ref_run.font.size = Pt(ref_font_size)
                    ref_run.font.name = font_name
                    _set_run_font(ref_run, font_name)

                _add_hyperlink(ref_para, f"https://doi.org/{doi}", f"doi:{doi}",
                               font_name, ref_font_size)

                if text_after_doi:
                    ref_run2 = ref_para.add_run(" " + text_after_doi)
                    ref_run2.font.size = Pt(ref_font_size)
                    ref_run2.font.name = font_name
                    _set_run_font(ref_run2, font_name)
            else:
                ref_run = ref_para.add_run(ref_text)
                ref_run.font.size = Pt(ref_font_size)
                ref_run.font.name = font_name
                _set_run_font(ref_run, font_name)

            # Hanging indent for references
            ref_para.paragraph_format.first_line_indent = Inches(-ref_hanging)
            ref_para.paragraph_format.left_indent = Inches(ref_hanging)
            ref_para.paragraph_format.space_after = Pt(2)
            ref_para.paragraph_format.space_before = Pt(1)
            _apply_line_spacing(ref_para, 1.0)

    # ── PAGE NUMBERS ──
    if journal_rules.get("page_numbers", True):
        _add_page_numbers(doc)

    # ── SAVE ──
    output_dir = OUTPUT_DIR / session_id
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = output_dir / "manuscript.docx"
    doc.save(str(docx_path))

    result = {"docx_path": str(docx_path), "pdf_path": None, "latex_path": None}

    # Try to convert to PDF
    pdf_path = _convert_to_pdf(docx_path, output_dir)
    if pdf_path:
        result["pdf_path"] = str(pdf_path)

    # Generate LaTeX
    try:
        latex_path = _generate_latex(structured_json, journal_rules, output_dir)
        if latex_path:
            result["latex_path"] = str(latex_path)
    except Exception:
        pass

    return result


# ─── FORMATTING HELPERS ───


def _setup_document_styles(doc, rules):
    """Configure document-wide default styles."""
    font_name = rules.get("font", "Times New Roman")
    font_size = rules.get("font_size", 12)

    # Set default style
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    _set_style_font(style, font_name)


def _apply_page_layout(doc, rules):
    """Apply page margins and layout from journal rules."""
    section = doc.sections[0]

    section.top_margin = _parse_margin(rules.get("margin_top", "1 inch"))
    section.bottom_margin = _parse_margin(rules.get("margin_bottom", "1 inch"))
    section.left_margin = _parse_margin(rules.get("margin_left", "1 inch"))
    section.right_margin = _parse_margin(rules.get("margin_right", "1 inch"))

    # Header/footer distance
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def _apply_two_column_layout(doc):
    """Apply two-column layout for IEEE/ACM style papers."""
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '480')  # Space between columns in twips (1/3 inch)
    cols.set(qn('w:equalWidth'), '1')


def _parse_margin(margin_str):
    """Parse margin string to docx measurement."""
    if isinstance(margin_str, (int, float)):
        return Inches(margin_str)

    margin_str = str(margin_str).lower().strip()
    nums = re.findall(r"[\d.]+", margin_str)
    if not nums:
        return Inches(1)

    val = float(nums[0])
    if "cm" in margin_str:
        return Cm(val)
    return Inches(val)


def _set_run_font(run, font_name):
    """Ensure font works for both ASCII and East Asian characters."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def _set_style_font(style, font_name):
    """Set font on a style object for cross-platform rendering."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def _apply_line_spacing(para, spacing_val):
    """Apply line spacing to a paragraph."""
    pf = para.paragraph_format
    if spacing_val == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif spacing_val == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif spacing_val == 2.0:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing = spacing_val


def _add_section_heading(doc, text, rules, level=1, section_num=None):
    """Add a properly formatted section heading with journal-specific numbering."""
    font_name = rules.get("font", "Times New Roman")
    font_size = rules.get("font_size", 12)
    heading_style = rules.get("heading_style", "bold_title_case")
    heading_size = rules.get("section_heading_size", font_size + 2)
    is_numbered = rules.get("section_numbered", False)

    heading_para = doc.add_paragraph()
    heading_para.space_before = Pt(18)
    heading_para.space_after = Pt(8)

    # Format heading text based on style
    if heading_style == "uppercase_roman":
        # IEEE-style: "I. INTRODUCTION", "II. RELATED WORK"
        formatted_text = text.upper()
        if is_numbered and section_num is not None and level == 1:
            roman = _to_roman(section_num)
            formatted_text = f"{roman}. {formatted_text}"
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif heading_style == "bold_title_case":
        formatted_text = text
        if is_numbered and section_num is not None and level == 1:
            formatted_text = f"{section_num}. {formatted_text}"
    else:
        formatted_text = text
        if is_numbered and section_num is not None and level == 1:
            formatted_text = f"{section_num}. {formatted_text}"

    heading_run = heading_para.add_run(formatted_text)
    heading_run.bold = rules.get("section_heading_bold", True)
    heading_run.font.name = font_name
    _set_run_font(heading_run, font_name)

    if level == 1:
        heading_run.font.size = Pt(heading_size)
        if heading_style != "uppercase_roman":
            heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_bottom_border(heading_para)
    elif level == 2:
        heading_run.font.size = Pt(heading_size - 1)
        heading_run.italic = True
    else:
        heading_run.font.size = Pt(font_size)
        heading_run.italic = True


def _add_bottom_border(para):
    """Add a bottom border line to a paragraph (for section headings)."""
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_full_section(doc, section, rules, line_spacing_val, section_num=None):
    """Add a complete section with heading and properly formatted content."""
    heading = section.get("heading", "")
    content = section.get("content", "")

    if heading:
        _add_section_heading(doc, heading, rules, level=1, section_num=section_num)

    if content:
        font_name = rules.get("font", "Times New Roman")
        font_size = rules.get("font_size", 12)
        para_indent = rules.get("paragraph_indent", 0.3)
        para_spacing = rules.get("paragraph_spacing_after", 4)

        # Split content into paragraphs
        if "\n\n" in content:
            paragraphs = content.split("\n\n")
        elif "\n" in content:
            lines = content.split("\n")
            paragraphs = []
            current = []
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    if current:
                        paragraphs.append(" ".join(current))
                        current = []
                else:
                    current.append(stripped)
            if current:
                paragraphs.append(" ".join(current))
        else:
            paragraphs = [content]

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if this looks like a sub-heading
            words = para_text.split()
            is_subheading = (
                len(words) <= 6
                and not para_text.endswith(".")
                and not para_text.endswith(",")
                and not para_text.endswith(":")
                and not any(c.isdigit() for c in para_text[:3])
                and para_text[0].isupper()
                and len(para_text) > 3
            )

            if is_subheading:
                sub_para = doc.add_paragraph()
                sub_para.space_before = Pt(14)
                sub_para.space_after = Pt(4)
                sub_run = sub_para.add_run(para_text)
                sub_run.bold = True
                sub_run.italic = True
                sub_run.font.size = Pt(font_size + 1)
                sub_run.font.name = font_name
                _set_run_font(sub_run, font_name)
            else:
                para = doc.add_paragraph()
                if para_indent > 0:
                    para.paragraph_format.first_line_indent = Inches(para_indent)
                para.space_after = Pt(para_spacing)
                _apply_line_spacing(para, line_spacing_val)

                run = para.add_run(para_text)
                run.font.size = Pt(font_size)
                run.font.name = font_name
                _set_run_font(run, font_name)


def _add_divider(doc):
    """Add a clean horizontal line divider between major sections."""
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.space_before = Pt(8)
    divider.space_after = Pt(8)
    pPr = divider._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _normalize_section_key(name: str) -> str:
    """Normalize section names for matching."""
    name = name.lower().strip()
    name = re.sub(r"^\d+\.?\s*", "", name)
    name = re.sub(r"[^\w\s]", "", name).strip()

    aliases = {
        "methods": "methods", "method": "methods",
        "methodology": "methods", "materials and methods": "materials and methods",
        "result": "results", "results": "results",
        "conclusion": "conclusion", "conclusions": "conclusion",
        "introduction": "introduction",
        "acknowledgments": "acknowledgments", "acknowledgements": "acknowledgments",
    }
    return aliases.get(name, name)


def _add_page_numbers(doc):
    """Add page numbers to the document footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_pre = para.add_run("— ")
    run_pre.font.size = Pt(9)
    run_pre.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar2)

    run_post = para.add_run(" —")
    run_post.font.size = Pt(9)
    run_post.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _convert_to_pdf(docx_path: Path, output_dir: Path) -> Optional[Path]:
    """Convert DOCX to PDF using pandoc if available."""
    pdf_path = output_dir / "manuscript.pdf"
    try:
        result = subprocess.run(
            ["pandoc", str(docx_path), "-o", str(pdf_path)],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(output_dir), str(docx_path)],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return None


def _add_hyperlink(paragraph, url, text, font_name, font_size):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _add_formatted_table(doc, table_data: dict, rules: dict, table_number: int = 1):
    """Add a professionally formatted table with auto-numbering."""
    font_name = rules.get("font", "Times New Roman")
    font_size = rules.get("font_size", 12)
    caption_pos = rules.get("table_caption_position", "above")

    caption = table_data.get("caption", "")
    if not caption or caption == f"Table {table_data.get('index', '')}":
        caption = f"Table {table_number}"
    elif not caption.lower().startswith("table"):
        caption = f"Table {table_number}. {caption}"

    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])

    if not headers and not rows:
        return

    total_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if total_cols == 0:
        return

    # Table caption
    cap_para = doc.add_paragraph()
    cap_para.space_before = Pt(16)
    cap_para.space_after = Pt(4)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_label = cap_para.add_run(caption)
    cap_label.bold = True
    cap_label.font.size = Pt(font_size - 1)
    cap_label.font.name = font_name
    _set_run_font(cap_label, font_name)

    # Create the table
    total_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_idx = 0

    # Header row
    if headers:
        for col_idx, header in enumerate(headers[:total_cols]):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(header or "")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(font_size - 1)
                    run.font.name = font_name
                    _set_run_font(run, font_name)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E8E8E8")
            cell._element.get_or_add_tcPr().append(shading)
        row_idx += 1

    # Data rows
    for data_row in rows:
        for col_idx, cell_val in enumerate(data_row[:total_cols]):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_val or "")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size - 1)
                    run.font.name = font_name
                    _set_run_font(run, font_name)
        row_idx += 1

    after_para = doc.add_paragraph()
    after_para.space_after = Pt(12)


def _add_image_to_doc(doc, image_data: dict, rules: dict, figure_number: int = 1):
    """Add an image to the document with auto-numbered figure caption."""
    font_name = rules.get("font", "Times New Roman")
    font_size = rules.get("font_size", 12)
    caption_pos = rules.get("figure_caption_position", "below")

    img_path = image_data.get("path", "")
    caption = image_data.get("caption", "")
    if not caption or caption == f"Figure {image_data.get('index', '')}":
        caption = f"Figure {figure_number}"
    elif not caption.lower().startswith("figure") and not caption.lower().startswith("fig"):
        caption = f"Figure {figure_number}. {caption}"

    if not img_path or not Path(img_path).exists():
        return

    try:
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.space_before = Pt(12)
        img_para.space_after = Pt(4)

        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(5))

        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.space_after = Pt(12)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(font_size - 1)
        cap_run.font.name = font_name
        _set_run_font(cap_run, font_name)

    except Exception:
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = placeholder.add_run(f"[{caption} — image not available]")
        run.italic = True
        run.font.size = Pt(font_size - 1)


# ─── LaTeX EXPORT ───


def _generate_latex(structured_json: dict, journal_rules: dict, output_dir: Path) -> Optional[Path]:
    """Generate a LaTeX file from structured content using journal-appropriate class."""
    citation_style = journal_rules.get("citation_style", "ieee")
    journal_name = journal_rules.get("journal_name", journal_rules.get("family_name", ""))

    # Determine document class
    if citation_style == "ieee" or "ieee" in journal_name.lower():
        doc_class = "IEEEtran"
        doc_options = "conference"
    elif "acm" in journal_name.lower():
        doc_class = "acmart"
        doc_options = "manuscript"
    elif "springer" in journal_name.lower():
        doc_class = "article"
        doc_options = "12pt"
    else:
        doc_class = "article"
        doc_options = "12pt, a4paper"

    title = _latex_escape(structured_json.get("title", "Untitled"))
    authors = structured_json.get("authors", [])
    abstract = _latex_escape(structured_json.get("abstract", ""))
    keywords = structured_json.get("keywords", [])
    sections = structured_json.get("sections", [])
    references = structured_json.get("references", [])

    # Build author string
    if isinstance(authors, list):
        author_str = " \\and ".join(_latex_escape(a) for a in authors)
    else:
        author_str = _latex_escape(str(authors))

    # Build LaTeX document
    lines = []
    lines.append(f"\\documentclass[{doc_options}]{{{doc_class}}}")
    lines.append("\\usepackage[utf8]{inputenc}")
    lines.append("\\usepackage{amsmath,amssymb}")
    lines.append("\\usepackage{graphicx}")
    lines.append("\\usepackage{hyperref}")
    lines.append("\\usepackage{url}")
    lines.append("")
    lines.append(f"\\title{{{title}}}")
    lines.append(f"\\author{{{author_str}}}")
    lines.append("")
    lines.append("\\begin{document}")
    lines.append("\\maketitle")
    lines.append("")

    # Abstract
    if abstract:
        lines.append("\\begin{abstract}")
        lines.append(abstract)
        lines.append("\\end{abstract}")
        lines.append("")

    # Keywords
    if keywords:
        kw_text = ", ".join(_latex_escape(k) for k in keywords)
        if doc_class == "IEEEtran":
            lines.append(f"\\begin{{IEEEkeywords}}")
            lines.append(kw_text)
            lines.append(f"\\end{{IEEEkeywords}}")
        else:
            lines.append(f"\\textbf{{Keywords:}} {kw_text}")
        lines.append("")

    # Sections
    for sec in sections:
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        if heading:
            lines.append(f"\\section{{{_latex_escape(heading)}}}")
        if content:
            lines.append(_latex_escape(content))
        lines.append("")

    # References
    if references:
        lines.append("\\begin{thebibliography}{99}")
        for i, ref in enumerate(references):
            ref_text = ref if isinstance(ref, str) else str(ref)
            lines.append(f"\\bibitem{{ref{i+1}}} {_latex_escape(ref_text)}")
        lines.append("\\end{thebibliography}")

    lines.append("")
    lines.append("\\end{document}")

    latex_content = "\n".join(lines)
    latex_path = output_dir / "manuscript.tex"
    latex_path.write_text(latex_content, encoding="utf-8")
    return latex_path


def _latex_escape(text: str) -> str:
    """Escape special LaTeX characters."""
    if not text:
        return ""
    replacements = {
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }
    # Don't escape backslashes that are already LaTeX commands
    text = text.replace('\\', r'\textbackslash{}')
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text
