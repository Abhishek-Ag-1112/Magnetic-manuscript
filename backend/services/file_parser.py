"""
File Parser Service — handles DOCX, PDF, TXT, and Markdown input.
Extracts text, tables, images, and attempts structural parsing.
Uses LLM heavily for complex documents where heuristics fail.
"""
import os
import re
import shutil
from pathlib import Path
from typing import Optional

try:
    import docx
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    docx = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import markdown
except ImportError:
    markdown = None

from utils.helpers import clean_text


# Only match VERY obvious standalone academic section headings
STRICT_SECTION_NAMES = {
    "abstract", "introduction", "background", "methods", "method",
    "methodology", "materials and methods", "results", "discussion",
    "conclusion", "conclusions", "references", "bibliography",
    "acknowledgments", "acknowledgements", "acknowledgment",
    "related work", "literature review", "experimental setup",
    "experimental results", "results and discussion", "data availability",
    "funding", "conflict of interest", "supplementary material",
    "star methods", "summary", "keywords", "significance",
    "author contributions", "data analysis",
}


def parse_file(file_path: str) -> dict:
    """Parse a manuscript file and return structured JSON."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".docx":
        return _parse_docx(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".txt":
        return _parse_txt(path)
    elif ext in (".md", ".markdown"):
        return _parse_markdown(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_docx(path: Path) -> dict:
    """Parse a DOCX file — extracts text, tables, and images."""
    if docx is None:
        raise ImportError("python-docx is required for DOCX parsing")

    doc = docx.Document(str(path))
    paragraphs = []
    raw_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""
        is_heading_style = "heading" in style_name or "title" in style_name

        # Analyze run-level formatting
        runs = para.runs
        all_bold = all(r.bold for r in runs if r.text.strip()) if runs else False
        has_any_text = any(r.text.strip() for r in runs)
        font_sizes = set()
        for r in runs:
            if r.text.strip() and r.font.size:
                font_sizes.add(r.font.size)

        is_structural_heading = False
        if is_heading_style:
            is_structural_heading = True
        elif all_bold and has_any_text and len(text.split()) <= 6:
            cleaned = re.sub(r"^\d+[\.)\]]\s*", "", text).strip().lower()
            if cleaned in STRICT_SECTION_NAMES:
                is_structural_heading = True

        paragraphs.append({
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "is_heading": is_structural_heading,
            "is_bold": all_bold,
            "font_sizes": font_sizes,
        })
        raw_lines.append(text)

    full_text = "\n".join(raw_lines)

    # Extract tables from DOCX
    extracted_tables = _extract_docx_tables(doc)

    # Extract images from DOCX
    extracted_images = _extract_docx_images(doc, path)

    # Check if we got meaningful structure
    heading_count = sum(1 for p in paragraphs if p["is_heading"])

    if heading_count >= 3:
        result = _extract_structure_from_styled(paragraphs, full_text)
    else:
        result = {
            "title": "",
            "authors": [],
            "abstract": "",
            "sections": [],
            "figures": [],
            "tables": [],
            "references": [],
            "raw_text": clean_text(full_text),
            "needs_llm_structuring": True,
        }

    # Merge extracted tables and images
    result["extracted_tables"] = extracted_tables
    result["extracted_images"] = extracted_images

    # Add table captions to the tables list
    for tbl in extracted_tables:
        caption = tbl.get("caption", "")
        if caption and caption not in result.get("tables", []):
            result.setdefault("tables", []).append(caption)

    return result


def _extract_docx_tables(doc) -> list:
    """Extract all tables from a DOCX document with their content."""
    tables = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            continue

        # Determine if first row is a header
        has_header = True

        # Build table dict
        table_dict = {
            "index": i + 1,
            "caption": f"Table {i + 1}",
            "headers": rows_data[0] if has_header else [],
            "rows": rows_data[1:] if has_header else rows_data,
            "total_rows": len(rows_data),
            "total_cols": len(rows_data[0]) if rows_data else 0,
        }

        # Try to find a caption from surrounding text
        # Look for "Table N" text in the caption
        for row in rows_data:
            for cell in row:
                tbl_match = re.match(r"(Table\s+\d+[.:]\s*.+)", cell, re.IGNORECASE)
                if tbl_match:
                    table_dict["caption"] = tbl_match.group(1)
                    break

        tables.append(table_dict)

    return tables


def _extract_docx_images(doc, docx_path: Path) -> list:
    """Extract images from a DOCX document."""
    images = []
    output_dir = docx_path.parent / "_images"

    try:
        # Access DOCX package relationships for images
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    content_type = image_part.content_type
                    ext = ".png"
                    if "jpeg" in content_type or "jpg" in content_type:
                        ext = ".jpg"
                    elif "gif" in content_type:
                        ext = ".gif"
                    elif "bmp" in content_type:
                        ext = ".bmp"

                    # Save image
                    output_dir.mkdir(parents=True, exist_ok=True)
                    img_filename = f"image_{len(images) + 1}{ext}"
                    img_path = output_dir / img_filename

                    with open(img_path, "wb") as f:
                        f.write(image_part.blob)

                    images.append({
                        "index": len(images) + 1,
                        "filename": img_filename,
                        "path": str(img_path),
                        "content_type": content_type,
                        "caption": f"Figure {len(images) + 1}",
                    })
                except Exception:
                    pass
    except Exception:
        pass

    return images


def _parse_pdf(path: Path) -> dict:
    """Parse a PDF file — extracts text and tables."""
    if pdfplumber is None:
        raise ImportError("pdfplumber is required for PDF parsing")

    raw_lines = []
    extracted_tables = []

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                raw_lines.append(page_text)

            # Extract tables from PDF
            try:
                page_tables = page.extract_tables()
                for tbl in page_tables:
                    if tbl and len(tbl) > 1:
                        extracted_tables.append({
                            "index": len(extracted_tables) + 1,
                            "caption": f"Table {len(extracted_tables) + 1} (page {page_num + 1})",
                            "headers": tbl[0] if tbl else [],
                            "rows": tbl[1:] if len(tbl) > 1 else [],
                            "total_rows": len(tbl),
                            "total_cols": len(tbl[0]) if tbl else 0,
                        })
            except Exception:
                pass

    full_text = "\n".join(raw_lines)

    return {
        "title": "",
        "authors": [],
        "abstract": "",
        "sections": [],
        "figures": [],
        "tables": [],
        "references": [],
        "raw_text": clean_text(full_text),
        "needs_llm_structuring": True,
        "extracted_tables": extracted_tables,
        "extracted_images": [],
    }


def _parse_txt(path: Path) -> dict:
    """Parse a plain text file."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    return {
        "title": "",
        "authors": [],
        "abstract": "",
        "sections": [],
        "figures": [],
        "tables": [],
        "references": [],
        "raw_text": clean_text(content),
        "needs_llm_structuring": True,
        "extracted_tables": [],
        "extracted_images": [],
    }


def _parse_markdown(path: Path) -> dict:
    """Parse a Markdown file — headings are explicit so this works well."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    lines = content.split("\n")
    paragraphs = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            paragraphs.append({
                "text": text,
                "style": f"Heading {level}",
                "is_heading": True,
                "is_bold": False,
                "font_sizes": set(),
            })
        else:
            paragraphs.append({
                "text": stripped,
                "style": "Normal",
                "is_heading": False,
                "is_bold": False,
                "font_sizes": set(),
            })

    full_text = content
    heading_count = sum(1 for p in paragraphs if p["is_heading"])

    if heading_count >= 3:
        result = _extract_structure_from_styled(paragraphs, full_text)
    else:
        result = {
            "title": "",
            "authors": [],
            "abstract": "",
            "sections": [],
            "figures": [],
            "tables": [],
            "references": [],
            "raw_text": clean_text(full_text),
            "needs_llm_structuring": True,
        }

    result.setdefault("extracted_tables", [])
    result.setdefault("extracted_images", [])
    return result


def _extract_structure_from_styled(paragraphs: list, full_text: str) -> dict:
    """Extract structure from documents that have proper heading styles."""
    result = {
        "title": "",
        "authors": [],
        "abstract": "",
        "sections": [],
        "figures": [],
        "tables": [],
        "references": [],
        "raw_text": clean_text(full_text),
        "needs_llm_structuring": False,
        "extracted_tables": [],
        "extracted_images": [],
    }

    current_section = None
    current_content = []
    found_title = False

    for para in paragraphs:
        text = para["text"]

        if not found_title and (para["is_heading"] or para.get("is_bold")):
            result["title"] = text
            found_title = True
            continue

        if para["is_heading"]:
            if current_section:
                content_text = "\n".join(current_content).strip()
                _store_section(result, current_section, content_text)

            current_section = text
            current_content = []
            continue

        current_content.append(text)

        # Detect figures and tables in text
        if re.match(r"^(Figure|Fig\.?)\s+\d+", text, re.IGNORECASE):
            result["figures"].append(text)
        elif re.match(r"^Table\s+\d+", text, re.IGNORECASE):
            result["tables"].append(text)

    if current_section and current_content:
        content_text = "\n".join(current_content).strip()
        _store_section(result, current_section, content_text)

    return result


def _store_section(result: dict, section_name: str, content: str):
    """Store a section into the appropriate field of the result dict."""
    lower_name = section_name.lower().strip()
    if lower_name == "abstract":
        result["abstract"] = content
    elif lower_name in ("references", "bibliography", "works cited"):
        result["references"] = _extract_references(content)
    elif lower_name == "keywords":
        keywords = [kw.strip() for kw in content.replace("\n", ",").split(",") if kw.strip()]
        result["keywords"] = keywords
    elif lower_name == "authors":
        authors = [a.strip() for a in content.replace("\n", ",").split(",") if a.strip()]
        result["authors"] = authors
    else:
        result["sections"].append({
            "heading": section_name,
            "content": content,
        })


def _extract_references(text: str) -> list:
    """Extract individual references from a reference block."""
    if not text:
        return []

    refs = re.split(r"\n\s*(?:\[\d+\]|\d+\.|\(\d+\))\s*", text)
    refs = [r.strip() for r in refs if r.strip()]

    if not refs:
        refs = [line.strip() for line in text.split("\n") if line.strip()]

    return refs
